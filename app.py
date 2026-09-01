import datetime
import os
import re
import io
import zipfile
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import parse_xml          # ⬅️ เพิ่มบรรทัดเดียว (ใช้ปั๊มสำเนาคู่ฉบับ)
from docxtpl import DocxTemplate
import gspread
import requests
import streamlit as st

# =========================================================================
# 📌 STREAMLIT PAGE CONFIG & EXECUTIVE UI STYLING
# =========================================================================
st.set_page_config(
    page_title="ระบบสร้างหนังสือราชการ (ASD)",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <link href="https://fonts.googleapis.com/css2?family=Prompt:wght@300;400;500;600;700&family=Sarabun:wght@400;600;700&display=swap" rel="stylesheet">
    <link href="https://fonts.googleapis.com/icon?family=Material+Icons" rel="stylesheet">

    <style>
        /* 1. ปรับแต่งฟอนต์และธีมพื้นหลังทั้งเว็บ */
        html, body, [class*="css"] {
            font-family: 'Prompt', 'Sarabun', sans-serif !important;
            background-color: #f8fafc !important;
        }

        /* 2. ซ่อน Header/Footer รกๆ ของ Streamlit Default */
        #MainMenu {visibility: hidden;}
        header {visibility: hidden;}
        footer {visibility: hidden;}

        /* 3. ตกแต่งหัวข้อหลัก (Executive Header Bar) */
        .main-header-box {
            background: linear-gradient(135deg, #0d47a1 0%, #1976d2 100%);
            padding: 28px 20px;
            border-radius: 16px;
            color: white;
            text-align: center;
            margin-bottom: 25px;
            box-shadow: 0 4px 20px rgba(13, 71, 161, 0.2);
        }
        .main-title {
            font-size: 1.8rem;
            font-weight: 700;
            margin: 0;
            color: #ffffff;
            letter-spacing: -0.3px;
        }
        .main-subtitle {
            font-size: 1.05rem;
            opacity: 0.9;
            margin-top: 6px;
            font-weight: 400;
        }

        /* 4. ปรับแต่งปุ่มกดหลัก (Primary Button) */
        div.stButton > button[kind="primary"] {
            background: linear-gradient(135deg, #1565c0 0%, #0d47a1 100%) !important;
            color: white !important;
            border-radius: 12px !important;
            padding: 12px 28px !important;
            font-size: 1.1rem !important;
            font-weight: 600 !important;
            border: none !important;
            box-shadow: 0 4px 14px rgba(13, 71, 161, 0.35) !important;
            transition: all 0.3s ease !important;
            width: 100% !important;
        }
        div.stButton > button[kind="primary"]:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 6px 20px rgba(13, 71, 161, 0.5) !important;
        }

        /* 5. ปรับแต่งปุ่มดาวน์โหลดเอกสาร (Download Buttons) */
        div.stDownloadButton > button {
            background: #ffffff !important;
            color: #2e7d32 !important;
            border: 2px solid #a5d6a7 !important;
            border-radius: 12px !important;
            font-weight: 600 !important;
            padding: 10px 20px !important;
            box-shadow: 0 2px 8px rgba(0,0,0,0.04) !important;
            transition: all 0.2s ease !important;
            width: 100% !important;
        }
        div.stDownloadButton > button:hover {
            background: #e8f5e9 !important;
            border-color: #2e7d32 !important;
            color: #1b5e20 !important;
        }

        /* 6. ตกแต่งส่วน Footer ท้ายเว็บ (แก้ไขให้อยู่กึ่งกลาง ไม่ให้ชื่อโดนตัดบรรทัด) */
        .custom-footer {
            font-family: 'Sarabun', sans-serif;
            text-align: center;
            background-color: #ffffff;
            border: 1px solid #e2e8f0;
            border-radius: 16px;
            padding: 35px 20px;
            margin: 40px auto 0 auto;
            max-width: 800px;
            color: #333333;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.04);
        }
        .footer-system-title {
            font-size: 1.05rem;
            font-weight: 700;
            color: #37474f;
            margin-bottom: 20px;
        }
        .footer-leader {
            margin-bottom: 18px;
            font-size: 0.95rem;
            color: #212121;
        }
        .advisor-row {
            display: flex;
            flex-wrap: wrap;
            justify-content: center;
            gap: 20px;
            margin: 15px 0 25px 0;
        }
        .advisor-col {
            flex: 1 1 300px;
            max-width: 350px;
            font-size: 0.9rem;
            color: #424242;
        }
        .dev-card {
            margin: 0 auto 20px auto;
            padding: 20px;
            border: 1px dashed #cfd8dc;
            border-radius: 12px;
            max-width: 500px;
            background-color: #ffffff;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04);
            text-align: center;
        }
        /* บังคับไม่ให้ชื่อเว้นบรรทัดเด็ดขาด */
        .custom-footer strong {
            white-space: nowrap;
        }
        .btn-report {
            display: inline-block;
            margin-top: 14px;
            padding: 8px 22px;
            background-color: #1565c0;
            color: #ffffff !important;
            border-radius: 20px;
            font-size: 0.9rem;
            font-weight: bold;
            text-decoration: none;
            box-shadow: 0 2px 5px rgba(21, 101, 192, 0.3);
            transition: all 0.2s ease-in-out;
        }
        .btn-report:hover {
            background-color: #0d47a1;
            box-shadow: 0 4px 10px rgba(13, 71, 161, 0.4);
        }
        .copyright-text {
            font-size: 0.85rem;
            color: #9e9e9e;
            margin-top: 25px;
        }
    </style>

    <div class="main-header-box">
        <div class="main-title">📄 ระบบสร้างหนังสือราชการอัตโนมัติ (ASD)</div>
        <div class="main-subtitle">ฝ่ายแจ้งความต้องการ กองพัสดุช่างอากาศ (ผคค.กพอ.ชอ.)</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# =========================================================================
# 📌 1. CONSTANTS & SPREADSHEET CONFIG (สมบูรณ์ 100% ห้ามตัดทอน)
# =========================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 📂 โฟลเดอร์ปลายทางสำหรับเก็บเอกสารที่สร้างเสร็จแล้ว
OUTPUT_DIR = os.path.join(BASE_DIR, "กห.ภายนอก บันทึกข้อความ")
if not os.path.exists(OUTPUT_DIR):
  os.makedirs(OUTPUT_DIR, exist_ok=True)

# 🔑 marker ล่องหน (zero-width space) ใช้ตรวจจับ "บล็อกรายการพัสดุ" ให้แม่น 100%
PART_MARK = "\u200b"

# 🖨️ namespace สำหรับสร้างกล่องข้อความลอย (ปั๊มสำเนาคู่ฉบับ)
VML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml"'
)

CREDENTIALS_DICT = {
    "type": "service_account",
    "project_id": "natural-choir-469013-f1",
    "private_key_id": "5e7f63e863ce2079edad42eb4c9054cc702403d3",
    "private_key": (
        "-----BEGIN PRIVATE"
        " KEY-----\nMIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQCba6rMpp8xvAIy\ndgrVrqK8NuEShuP89VZ4cGEt03nTFibSW8wv0OlhQ9gHshmsyJ19CaYTbGMFk+FN\niLqU4PHFYoQp681VdGY8XbuR0HpiR+3zUqxv6Ps49LaDKi9aj+jwhlCkS8nrHr/D\neFojhyUCf88qNE2/EKy7oHxtL7t1siI56gXDl400qsRYDvEipWcng7868C7Gq4qB\nQGuyj5VqUbkmpFrJ3z/IHVCph3gR5DqSTLF7boTVWpq57iBvdwB1ti9Dh6yqteHP\neQNEKFiq6ImXTmUj1cUvydqK3VJp0kmVEqX3FiLVPp8gH8r6opH4iQET0QKEcpnu\nrQin/0LrAgMBAAECggEACRwVICG0EF+LZsHHKMyE+rCMyLU3WnWsN9WpHIfaCtFa\n59bm76E4d1xZUTooicQd1we1wEsIvXme1G3rNxwS5RWSgZKHWow+tyhsp238073m\nPe4z7+tMA66e9APeUASXYWG+CviWQT3FfQOU5Pg1kcvurJ6cNVLDjSEpqeIJBCuc\nC8qU3/z/KrzgvRWRMuZjyPsWltDJ2EwUxRtqTiFmYL6933sOUGTAsS7V6az0T1JI\nl9NTrOnCuGuVLZq7dYAoAjUihtfG0NlbWciagQTL4bRpCP98bCIBsAenSzbeFFbk\nJjGIhPwISlx5nGaVrJMmJYTpKY93ErfoGYGF6g31oQKBgQDIMOeFPXOG9z9w6Ddo\nS35Py0JdEaRDkweD1mVkkp5nSzV/WV3zFL43scHlLaszg9c7Bz/LvyU8gGoxokqx\nAp8jH9dBoFrsHTqejY7VrBMuKNWENWaTYOtQ8sA7N2pcfNtQZBlRQDo/e3U+piQW\nkZUU2hY4RQIUbBT5xodG8Lh37wKBgQDGv52akNtOUkxCsI38UyMXZv5TBHZCkNA0\ny12Ctk8/N496V6doRwHb4n7YeoOiA+GkJTly03irFd/Gwim3vXaNpmOLXKVPkoju\n/Q2GQVCkE7FzNLseNGdEYs0+dRIyZubyRuYLrtMvE3LreyAWTJAx/Uc+M8xPmsnY\nOJEixuyIxQKBgDlD3AZ+NJzn/yrSEn9wEPrMXvh2gnGeDmlFHA3v7wYHOo9qRfix\n91PBMoDXVoDO9vN4uGQVEpbC+R2nmgwWfuUyR4YLU9b06X7PaYtvxLDQl3tRNz1z\nXPzz489Mo80/HhFaAPAAGmlsbHZ2Wh5mmKm1VOPVwamL3Vgx1SKS24HbAoGBAJ9T\n2Uqhuc2t2AjWdNzE4SrPnC59Mzjl0qOgLFSvRhRNvC29uyyzT+AwULPGc2QcbHUk\nikttEB1HKd+yo7Lypemka8S6/qMtu6yrHH52OelvCCBtM1xhci+2bQcW3wGc0KOF\nBsJy4kWo98WjLPPzaN1KSCSrbaybUBiQMHmKsvBpAoGAGlY5BU5H0ffV6ynen4s8\nxObPoO1akCeplu75mo0NRn+X48ISQjp86MrlMn7cY/DcobiBUOLUWGrQveDpVsrc\nO+fBF5obJZi4fRMF2RwGI47nk1ubcXtZs26EtzDwJyNygbOZiEMVWKliNuc5Ci6B\nQ3QjHPocAsYJSI7beZGhf0M=\n-----END"
        " PRIVATE KEY-----\n"
    ),
    "client_email": "printlog@natural-choir-469013-f1.iam.gserviceaccount.com",
    "client_id": "101691942194098646397",
    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
    "token_uri": "https://oauth2.googleapis.com/token",
    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
    "client_x509_cert_url": (
        "https://www.googleapis.com/robot/v1/metadata/x509/printlog%40natural-choir-469013-f1.iam.gserviceaccount.com"
    ),
    "universe_domain": "googleapis.com",
}

SPREADSHEET_ID = "1icuJgU9n75mUs5n2KlttxQihRYO6hMayeQHkwL2J5do"
UPLOADEDFILES2_ID = "1X0NA8KDxpXGiCEpEkLZiLfpp28myR-uBnUGB9Z0KENk"

LETTER_NORMAL_URL = "https://1drv.ms/w/c/3aea505ab7e3838e/IQD4uB7vZVuAQ6ker632o8phAZV3DZUeg_lHTT8Uft5MVJY?download=1"
LETTER_LP9_URL    = "https://1drv.ms/w/c/3aea505ab7e3838e/IQBYVIK04q5zRbj5WvK_nFUzAY-Vvax8oN7qbXLrA4xDbxs?download=1"
MEMO_NORMAL_URL   = "https://1drv.ms/w/c/3aea505ab7e3838e/IQDVF-keiaNyQpdfzT10kDMjAewJUK9KILNAF9IAzrt6TQg?download=1"
MEMO_LP9_URL      = "https://1drv.ms/w/c/3aea505ab7e3838e/IQC2dfosh_q5SIxtTlJ90RMaAXXt7qMBTZ5wjRs_lR6uNs0?download=1"


# =========================================================================
# 📌 2. FULL CONSTANT MAPPINGS (ครบถ้วนสมบูรณ์ 100%)
# =========================================================================
FULL_HELICOPTER_NAMES = {
    "S92A": "เฮลิคอปเตอร์แบบที่ ๑๐ (S-92A)",
    "BELL": "เฮลิคอปเตอร์แบบที่ ๖ ง (BELL-412EP)",
    "BELL-412EP": "เฮลิคอปเตอร์พระราชพาหนะแบบที่ ๖ ง. (BELL-412EP)",
    "H225M": "เฮลิคอปเตอร์แบบที่ ๑๑ (H225M)",
    "S70i": "เฮลิคอปเตอร์แบบที่ ๑๒ (S-70i)",
    "H135": "เฮลิคอปเตอร์แบบที่ ๑๓ (H-135)",
    "TEST": "เฮลิคอปเตอร์แบบที่ ๐๐",
    "F16": "เครื่องบินขับไล่แบบที่ ๑๙/ก (F-16A/B)",
    "F5": "เครื่องบินขับไล่แบบที่ ๑๘ ข/ค (F-5E/F)",
    "AU23A": "เครื่องบินโจมตีธุรการแบบที่ ๒ (AU-23A)",
    "ATR": "เครื่องบินลำเลียงแบบที่ ๑๖ ก (ATR72-500/600)",
    "T6C": "เครื่องบินฝึกแบบที่ ๒๒ (T-6C)",
    "AT6": "เครื่องบินโจมตีแบบที่ ๘ (AT-6TH)",
    "ALPHA_JET": "เครื่องบินโจมตีแบบที่ ๗ (Alpha Jet)",
    "BT67": "เครื่องบินลำเลียงแบบที่ ๒ ก (BT-67)",
    "C130": "เครื่องบินลำเลียงแบบที่ ๘ (C-130H/H-30)",
    "AT0": "เครื่องบินลำเลียงแบบที่ ๑๖ ก",
}

SHORT_HELICOPTER_NAMES = {
    "S92A": "ฮ.๑๐ (S-92A)",
    "BELL": "ฮ.๖ ง (BELL-412EP)",
    "BELL-412EP": "ฮ.พระราชพาหนะแบบที่ ๖ ง. (BELL-412EP)",
    "H225M": "ฮ.๑๑ (H225M)",
    "S70i": "ฮ.๑๒ (S-70i)",
    "H135": "ฮ.๑๓ (H-135)",
    "TEST": "ฮ.๐๗ (TEST)",
    "F16": "บ.ข.๑๙/ก (F-16A/B)",
    "F5": "บ.ข.๑๘ ข/ค (F-5E/F)",
    "AU23A": "บ.จธ.๒ (AU-23A)",
    "ATR": "บ.ล.๑๖ ก (ATR72-500/600)",
    "T6C": "บ.ฝ.๒๒ (T-6C)",
    "AT6": "บ.จ.๘ (AT-6TH)",
    "ALPHA_JET": "บ.จ.๗ (Alpha Jet)",
    "BT67": "บ.ล.๒ ก (BT-67)",
    "C130": "บ.ล.๘ (C-130H/H-30)",
    "AT0": "บ.ล.๑๖ ก (ATR72-500/600)",
}

HELICOPTER_CONTRACT_MAP = {
    "S92A": "๒๑/๒๕๖๙ จอ.",
    "BELL": "๘/๒๕๖๙ จอ.",
    "BELL-412EP": "๑๘/๒๕๖๙ จอ.",
    "H225M": "๑๓/๒๕๖๙ จอ.",
    "S70i": "๑๙/๒๕๖๙ จอ.",
    "H135": "๑๕/๒๕๖๘ จอ.",
    "TEST": "๐๑/๒๕๖๘ จอ.",
    "F16": "๒/๒๕๖๙ จอ.",
    "F5": "๑/๒๕๖๙ จอ.",
    "AU23A": "๑๒/๒๕๖๙ จอ.",
    "ATR": "๑๙/๒๕๖๙ จอ.",
    "T6C": "๔/๒๕๖๙ จอ.",
    "AT6": "๔/๒๕๖๙ จอ.",
    "ALPHA_JET": "๓/๒๕๖๙ จอ.",
    "BT67": "๓๘/๒๕๖๙ จอ.",
    "C130": "๓๙/๒๕๖๙ จอ.",
    "AT0": "๑๙/๒๕๖๙ จอ.",
}

UNIT_MAPPING = {
    "EA": "EACH",
    "KT": "KIT",
    "RL": "ROLL",
    "RO": "ROLL",
    "BX": "BOX",
    "GL": "GALLON",
    "CN": "CAN",
    "PR": "PAIR",
    "PC": "PIECE",
    "PCS": "PIECES",
    "ST": "SET",
}

THAI_MONTHS_FULL = [
    "มกราคม",
    "กุมภาพันธ์",
    "มีนาคม",
    "เมษายน",
    "พฤษภาคม",
    "มิถุนายน",
    "กรกฎาคม",
    "สิงหาคม",
    "กันยายน",
    "ตุลาคม",
    "พฤศจิกายน",
    "ธันวาคม",
]
THAI_MONTHS_SHORT = [
    "ม.ค.",
    "ก.พ.",
    "มี.ค.",
    "เม.ย.",
    "พ.ค.",
    "มิ.ย.",
    "ก.ค.",
    "ส.ค.",
    "ก.ย.",
    "ต.ค.",
    "พ.ย.",
    "ธ.ค.",
]


def to_thai_num(num):
  if num is None:
    return ""
  thai_digits = {
      "0": "๐",
      "1": "๑",
      "2": "๒",
      "3": "๓",
      "4": "๔",
      "5": "๕",
      "6": "๖",
      "7": "๗",
      "8": "๘",
      "9": "๙",
  }
  return "".join(thai_digits.get(ch, ch) for ch in str(num))


# =========================================================================
# 🖨️ ฟังก์ชันปั๊ม "สำเนาคู่ฉบับ" กลางบนสุดของหน้ากระดาษ (ลอย ไม่ดันเนื้อหา)
# =========================================================================
def stamp_copy_label(
    doc,
    text="สำเนาคู่ฉบับ",
    font="TH SarabunPSK",
    size_pt=18,
    top_pt=20,          # ระยะจากขอบบนกระดาษ (ปรับเลขนี้ทีละ 4 ถ้าอยากขยับขึ้น/ลง)
    width_pt=240,
    color="000000",     # "FF0000" = แดง
):
  """
  วางข้อความลอยแบบ absolute อ้างอิง 'ขอบกระดาษ' (page)
  → กึ่งกลางบนสุดเหนือตราครุฑ และไม่ดันเนื้อหาลงแม้แต่บรรทัดเดียว
  """
  if not doc.paragraphs:
    return False

  anchor = doc.paragraphs[0]
  sz = int(size_pt * 2)  # half-point
  xml = (
      f'<w:r {VML_NS}><w:pict>'
      f'<v:rect id="CopyStampBox" style="position:absolute;'
      f'margin-left:0;margin-top:{top_pt}pt;'
      f'width:{width_pt}pt;height:{size_pt + 8}pt;'
      f'mso-position-horizontal:center;'
      f'mso-position-horizontal-relative:page;'
      f'mso-position-vertical-relative:page;'
      f'z-index:251659264" filled="f" stroked="f">'
      f'<v:textbox inset="0,0,0,0"><w:txbxContent><w:p>'
      f'<w:pPr><w:jc w:val="center"/>'
      f'<w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
      f'<w:r><w:rPr>'
      f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
      f'<w:b/><w:bCs/><w:color w:val="{color}"/>'
      f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/></w:rPr>'
      f'<w:t xml:space="preserve">{text}</w:t>'
      f'</w:r></w:p></w:txbxContent></v:textbox>'
      f'</v:rect></w:pict></w:r>'
  )
  anchor._p.append(parse_xml(xml))
  return True


# =========================================================================
# 🚀 3. LOGIC ประมวลผลสร้างเอกสาร (รักษา Logic เดิม 100% ห้ามตัดทอน)
# =========================================================================
def generate_documents_process(
    unique_id, template_source_mode, log_func, finish_callback
):
  try:
    log_func(f"🚀 เริ่มต้นสร้างเอกสารสำหรับ: {unique_id}")
    upper_id = unique_id.strip().upper()
    is_repair_job = "LP9" in upper_id or "LP 9" in upper_id

    out_letter_path = os.path.join(OUTPUT_DIR, f"หนังสือภายนอก_{unique_id}.docx")
    out_memo_path = os.path.join(OUTPUT_DIR, f"บันทึกข้อความ_{unique_id}.docx")

    if is_repair_job:
      target_letter_url = LETTER_LP9_URL
      target_memo_url = MEMO_LP9_URL
      local_letter_file = "template_letter_lp9.docx"
      local_memo_file = "template_memo_lp9.docx"
    else:
      target_letter_url = LETTER_NORMAL_URL
      target_memo_url = MEMO_NORMAL_URL
      local_letter_file = "template_letter_normal.docx"
      local_memo_file = "template_memo_normal.docx"

    path_letter = os.path.join(BASE_DIR, local_letter_file)
    path_memo = os.path.join(BASE_DIR, local_memo_file)

    if template_source_mode == "online":
      log_func("⏳ กำลังดาวน์โหลด Template Word จาก OneDrive...")
      open(path_letter, "wb").write(
          requests.get(target_letter_url, allow_redirects=True).content
      )
      open(path_memo, "wb").write(
          requests.get(target_memo_url, allow_redirects=True).content
      )
    else:
      log_func("📁 กำลังดึงไฟล์แม่แบบจากภายในเครื่อง...")
      if not os.path.exists(path_letter) or not os.path.exists(path_memo):
        raise FileNotFoundError(
            f"ไม่พบไฟล์แม่แบบ {local_letter_file} หรือ {local_memo_file} ในเครื่อง"
        )

    if "BELL-412EP" in upper_id or "412EP" in upper_id:
      raw_model = "BELL-412EP"
    else:
      parts = unique_id.split("-")
      raw_model = parts[1].strip() if len(parts) >= 3 else unique_id

    full_model = FULL_HELICOPTER_NAMES.get(raw_model, raw_model).strip()
    short_model = SHORT_HELICOPTER_NAMES.get(raw_model, raw_model).strip()
    contract_no = HELICOPTER_CONTRACT_MAP.get(raw_model, "-").strip()

    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    gc = gspread.service_account_from_dict(CREDENTIALS_DICT, scopes=scopes)

    ss = gc.open_by_key(SPREADSHEET_ID)
    logs_data = ss.worksheet("PrintLogs").get_all_values()

    target_sheet_name = ""
    for i in range(len(logs_data) - 1, 0, -1):
      if len(logs_data[i]) > 1 and logs_data[i][1].strip().upper() == upper_id:
        if len(logs_data[i]) > 3:
          target_sheet_name = logs_data[i][3].strip()
          break

    if not target_sheet_name:
      raise ValueError(f"ไม่พบประวัติรหัส {unique_id} ในชีต PrintLogs")

    master_data = (
        gc.open_by_key(UPLOADEDFILES2_ID)
        .worksheet(target_sheet_name)
        .get_all_values()
    )
    if len(master_data) <= 3:
      raise ValueError("ข้อมูลพัสดุในชีตน้อยเกินไป (ไม่พบหัวตาราง)")

    master_headers = [str(h).upper().strip() for h in master_data[3]]
    idx_item = next(
        (
            i
            for i, h in enumerate(master_headers)
            if "ลำดับ" in h or "ITEM" in h or "NO" in h
        ),
        0,
    )
    idx_pn = next(
        (i for i, h in enumerate(master_headers) if "P/N" in h or "PN" in h), -1
    )
    idx_nsn = next(
        (i for i, h in enumerate(master_headers) if "NSN" in h), -1
    )
    idx_sn = next(
        (i for i, h in enumerate(master_headers) if "S/N" in h or "SN" in h), -1
    )
    idx_name = next(
        (
            i
            for i, h in enumerate(master_headers)
            if "NAME" in h or "DESCRIPTION" in h or "รายการ" in h
        ),
        -1,
    )
    idx_qty = next(
        (
            i
            for i, h in enumerate(master_headers)
            if "QTY" in h or "QUANTITY" in h
        ),
        -1,
    )
    idx_ua = next(
        (
            i
            for i, h in enumerate(master_headers)
            if "UI" in h
            or "U/I" in h
            or "UA" in h
            or "UNIT" in h
            or "หน่วย" in h
        ),
        -1,
    )
    idx_price = next(
        (i for i, h in enumerate(master_headers) if "ราคา" in h), -1
    )
    idx_remark = next(
        (
            i
            for i, h in enumerate(master_headers)
            if "หมายเหตุ" in h or "REMARK" in h
        ),
        -1,
    )

    valid_parts = []
    for i in range(4, len(master_data)):
      row = master_data[i]
      if len(row) <= max(idx_item, idx_pn, idx_name):
        continue
      pn_val = row[idx_pn].strip() if idx_pn != -1 and len(row) > idx_pn else ""
      name_val = (
          row[idx_name].strip()
          if idx_name != -1 and len(row) > idx_name
          else ""
      )
      if not row[idx_item] and not pn_val and not name_val:
        break
      if "หมายเหตุ" in str(name_val) or "หมายเหตุ" in str(row[idx_item]):
        break

      if pn_val or name_val:
        price_val = (
            row[idx_price].strip().upper()
            if idx_price != -1 and len(row) > idx_price
            else ""
        )
        remark_val = (
            row[idx_remark].strip().upper()
            if idx_remark != -1 and len(row) > idx_remark
            else ""
        )

        is_explicit_no_bid = (
            "NO BID" in price_val
            or "NOBID" in price_val
            or price_val == "0"
            or "NO BID" in remark_val
            or "NOBID" in remark_val
            or "ไม่เสนอราคา" in remark_val
        )

        if not is_explicit_no_bid:
          raw_ua = "EA"
          if idx_ua != -1 and len(row) > idx_ua:
            cell_text = row[idx_ua].strip().upper()
            if cell_text:
              raw_ua = cell_text
          final_ua = UNIT_MAPPING.get(raw_ua, raw_ua)

          valid_parts.append({
              "pn": (
                  row[idx_pn].strip()
                  if idx_pn != -1 and len(row) > idx_pn
                  else "-"
              ),
              "nsn": (
                  row[idx_nsn].strip()
                  if idx_nsn != -1 and len(row) > idx_nsn
                  else "-"
              ),
              "sn": (
                  row[idx_sn].strip()
                  if idx_sn != -1 and len(row) > idx_sn
                  else "-"
              ),
              "name": (
                  row[idx_name].strip()
                  if idx_name != -1 and len(row) > idx_name
                  else "-"
              ),
              "qty": (
                  row[idx_qty].strip()
                  if idx_qty != -1 and len(row) > idx_qty
                  else "1"
              ),
              "ua": final_ua,
          })

    if not valid_parts:
      valid_parts.append({
          "pn": "-",
          "nsn": "-",
          "sn": "-",
          "name": f"พัสดุตามใบแจ้ง {unique_id}",
          "qty": "1",
          "ua": "รายการ",
      })

    formatted_parts = [
        {"itemNoThai": to_thai_num(i + 1), **item}
        for i, item in enumerate(valid_parts)
    ]
    total_count = len(formatted_parts)
    item_count_thai = to_thai_num(total_count)

    if total_count == 1:
      p = formatted_parts[0]
      clean_name = re.sub(r"\s+", " ", str(p["name"])).strip()
      code_ref = f"S/N {p['sn']}" if is_repair_job else f"NSN {p['nsn']}"
      part_details_letter = (
          f"คือ {clean_name} P/N {p['pn']} {code_ref} จำนวน {p['qty']}"
          f" {p['ua']}"
      )
      part_details_memo = part_details_letter
    elif 2 <= total_count <= 3:
      list_letter, list_memo = [], []
      for index, p in enumerate(formatted_parts):
        clean_name = re.sub(r"\s+", " ", str(p["name"])).strip()
        code_ref = f"S/N {p['sn']}" if is_repair_job else f"NSN {p['nsn']}"
        sub_item_thai = to_thai_num(index + 1)

        str_letter = (
            f"\n---------------{p['itemNoThai']}. {clean_name} P/N {p['pn']}"
            f" {code_ref} จำนวน {p['qty']} {p['ua']}"
        )
        # ✅ FIX ①-A : ใส่ marker ล่องหนไว้ท้ายช่องว่าง (มองไม่เห็นในเอกสาร)
        list_letter.append(
            str_letter.replace("---------------", "             " + PART_MARK)
        )

        # 📌 โค้ดต้นฉบับของคุณ 100% (ฝั่ง memo มี \u200b เป็น marker อยู่แล้ว)
        list_memo.append(
            f"                      ๑.{sub_item_thai}\u200b"
            f" \u200b{clean_name} P/N {p['pn']} {code_ref} จำนวน {p['qty']}"
            f" {p['ua']}"
        )

      # 📌 แยก "ดังนี้:" ด้วย \a เพื่อเว้นล่าง 6 PT และเชื่อม 1.1, 1.2 ให้ติดกันในย่อหน้าเดียวด้วย \t\n เหมือนต้นฉบับ!
      part_details_letter = "ดังนี้:\a" + "".join(list_letter).lstrip("\n")
      part_details_memo = "ดังนี้:\a" + "\n".join(list_memo).lstrip("\n")

    else:
      part_details_letter = f"รายละเอียดตามใบแจ้งความต้องการเลขที่ {unique_id}"
      part_details_memo = part_details_letter

    now = datetime.datetime.now()
    year_thai = now.year + 543
    short_year_thai = to_thai_num(str(year_thai)[-2:])
    full_date = f"{THAI_MONTHS_FULL[now.month - 1]} {to_thai_num(year_thai)}"
    short_date = f"{THAI_MONTHS_SHORT[now.month - 1]}{short_year_thai}"

    doc_letter = DocxTemplate(path_letter)
    doc_letter.render({
        "FULL_DATE": full_date,
        "FULL_MODEL_NAME": full_model,
        "CONTRACT_MAP_NO": contract_no,
        "ITEM_COUNT_THAI": item_count_thai,
        "PART_DETAILS_STRING": part_details_letter,
    })
    doc_letter.save(out_letter_path)

    doc_memo = DocxTemplate(path_memo)
    doc_memo.render({
        "FULL_DATE": full_date,
        "SHORT_DATE": short_date,
        "SHORT_MODEL_THAI": short_model,
        "CONTRACT_REF": contract_no,
        "ITEM_COUNT_THAI": item_count_thai,
        "PART_DETAILS_STRING": part_details_memo,
        "UNIQUE_ID": unique_id,
    })
    doc_memo.save(out_memo_path)

    # =========================================================================
    # 🎯 7. จัดย่อหน้าและระยะห่าง (คงลอจิกต้นฉบับ + FIX ① ระยะบล็อกพัสดุ)
    # =========================================================================
    list_pattern = re.compile(r"^([๑-๙0-9]+)\.(?:([๑-๙0-9]+))?\s")

    # === จัดการหนังสือภายนอก (Letter) ===
    doc_l = docx.Document(out_letter_path)
    prev_was_list_l = False
    prev_was_sub_l = False

    for p in doc_l.paragraphs:
        raw_text = p.text
        text = raw_text.strip()
        if not text:
            continue

        # ✅ FIX ①-B : บล็อกรายการพัสดุ → คุมด้วย space_before ตัวเดียว ไม่บวกทบ
        if PART_MARK in raw_text:
            pf = p.paragraph_format
            pf.line_spacing = 1.0
            pf.space_before = Pt(6)     # ระยะจาก "ดังนี้:" ลงมา 6 PT (ครั้งเดียว)
            pf.space_after = Pt(0)
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.keep_together = True     # ไม่ให้พัสดุขาดหน้ากลางคัน
            prev_was_list_l = True
            prev_was_sub_l = True
            continue

        # 📌 ล็อกตราครุฑและโครงสร้าง Template ไม่ให้เด้ง: ปรับระยะบรรทัด 0.90 เฉพาะส่วนที่เป็นเนื้อหาหลักเท่านั้น
        is_body_text = False
        if (
            list_pattern.match(text) or
            text.startswith("ตามอ้างถึง") or
            text.startswith("จึงขอให้") or
            text.startswith("จึงเรียนมา") or
            text.startswith("ด้วย") or
            text.startswith("รายละเอียดตามใบแจ้ง") or
            "ดังนี้:" in text or
            re.match(r"^[๑-๙]\.[๑-๙]", text)
        ):
            is_body_text = True

        if is_body_text:
            p.paragraph_format.line_spacing = 1.0

        # ✅ FIX ①-C : "ดังนี้:" ไม่ใส่ space_after แล้ว (ย้ายไปคุมที่ space_before ของบล็อกพัสดุ)
        p.paragraph_format.space_after = Pt(0)
        if text.endswith("ดังนี้:"):
            p.paragraph_format.keep_with_next = True

        match = list_pattern.match(text)
        if match:
            is_sub = match.group(2) is not None
            if not prev_was_list_l:
                p.paragraph_format.space_before = Pt(6) 
            else:
                if not is_sub and prev_was_sub_l:
                    p.paragraph_format.space_before = Pt(6) 
                else:
                    p.paragraph_format.space_before = Pt(0)
            
            prev_was_list_l = True
            prev_was_sub_l = is_sub
        else:
            if prev_was_list_l:
                p.paragraph_format.space_before = Pt(6) 
            prev_was_list_l = False
            prev_was_sub_l = False

        if (
            text.startswith("ตามอ้างถึง")
            or text.startswith("จึงขอให้")
            or text.startswith("จึงเรียนมา")
        ):
            p.paragraph_format.first_line_indent = Pt(45)
            p.paragraph_format.left_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    doc_l.save(out_letter_path)

    # === จัดการบันทึกข้อความ (Memo / หนังสือภายใน) ===
    doc_m = docx.Document(out_memo_path)
    prev_was_list_m = False
    prev_was_sub_m = False

    for p in doc_m.paragraphs:
        raw_text = p.text
        text = raw_text.strip()
        if not text:
            continue

        # ✅ FIX ①-B : บล็อกรายการพัสดุฝั่งบันทึกข้อความ
        if PART_MARK in raw_text:
            pf = p.paragraph_format
            pf.line_spacing = 1.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(0)
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.keep_together = True
            p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
            prev_was_list_m = True
            prev_was_sub_m = True
            continue

        # 📌 ปรับระยะบรรทัด 0.90 เฉพาะส่วนที่เป็นเนื้อหาหลัก
        is_body_m = False
        if (
            list_pattern.match(text) or
            text.startswith("จึงเรียนมา") or
            text.startswith("เพื่อลงชื่อ") or
            text.startswith("เพื่อโปรด") or
            text.startswith("ด้วย") or
            text.startswith("ตามอ้างถึง") or
            text.startswith("รายละเอียดตามใบแจ้ง") or
            "ดังนี้:" in text or
            re.match(r"^[๑-๙]\.[๑-๙]", text)
        ):
            is_body_m = True

        if is_body_m:
            p.paragraph_format.line_spacing = 1.0

        # ✅ FIX ①-C : ปิด space_after ทั้งหมด กันบวกทบ
        p.paragraph_format.space_after = Pt(0)
        if text.endswith("ดังนี้:"):
            p.paragraph_format.keep_with_next = True

        match = list_pattern.match(text)
        if match:
            is_sub = match.group(2) is not None
            if not prev_was_list_m:
                p.paragraph_format.space_before = Pt(6)
            else:
                if not is_sub and prev_was_sub_m:
                    p.paragraph_format.space_before = Pt(6)
                else:
                    p.paragraph_format.space_before = Pt(0)
                    
            prev_was_list_m = True
            prev_was_sub_m = is_sub
        else:
            if prev_was_list_m:
                p.paragraph_format.space_before = Pt(6)
            prev_was_list_m = False
            prev_was_sub_m = False

        if (
            re.match(r"^[๑-๙]\.\s", text)
            or text.startswith("จึงเรียนมา")
            or text.startswith("เพื่อลงชื่อ")
            or text.startswith("เพื่อโปรด")
        ):
            p.paragraph_format.first_line_indent = Pt(66)
            p.paragraph_format.left_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
            
        elif re.match(r"^[๑-๙]\.[๑-๙]", text):
            p.paragraph_format.first_line_indent = Pt(72)
            p.paragraph_format.left_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    doc_m.save(out_memo_path)

    # =========================================================================
    # 📌 โค้ดส่วนสร้างสำเนาคู่ฉบับ (สร้างสำหรับทั้งหนังสือภายนอกและบันทึกข้อความ)
    # =========================================================================
    out_copy_letter_path = os.path.join(OUTPUT_DIR, f"สำเนาคู่ฉบับ_กห.ภายนอก_{unique_id}.docx")
    out_copy_memo_path = os.path.join(OUTPUT_DIR, f"สำเนาคู่ฉบับ_บันทึกข้อความ_{unique_id}.docx")

    def make_copy_doc(src_path, dest_path, doc_type):
        doc_copy = docx.Document(src_path)

        # ✅ FIX ② : ปั๊ม "สำเนาคู่ฉบับ" กลางบนสุดเหนือตราครุฑ แบบลอย ไม่ดันเนื้อหาลงเลย
        stamp_copy_label(doc_copy, "สำเนาคู่ฉบับ", size_pt=32, top_pt=18)

        # 1. ค้นหาจุดตัดเพื่อลบลายเซ็นด้านล่าง
        delete_start_idx = -1
        for i, p in enumerate(doc_copy.paragraphs):
            text_no_space = p.text.replace(" ", "").replace("\u200b", "")
            # ของบันทึกข้อความ: ลบตั้งแต่ใต้คำว่า หน.ผคค.กพอ.ชอ.
            if doc_type == "memo" and "หน.ผคค.กพอ.ชอ." in text_no_space:
                delete_start_idx = i + 1 
                break
            # ของหนังสือภายนอก: ลบตั้งแต่ใต้คำว่า เจ้ากรมช่างอากาศ
            elif doc_type == "letter" and "เจ้ากรมช่างอากาศ" in text_no_space:
                delete_start_idx = i + 1
                break

        # 2. หั่นย่อหน้าตั้งแต่จุดตัดลงไปทิ้งทั้งหมด
        if delete_start_idx != -1 and delete_start_idx < len(doc_copy.paragraphs):
            for p in doc_copy.paragraphs[delete_start_idx:]:
                p_element = p._element
                p_element.getparent().remove(p_element)
                p._element, p._p = None, None

        # 3. เติมบล็อก "ร่าง พิมพ์ ทาน" แบบชิดขวาสุด
        for _ in range(4):  # เคาะบรรทัดว่างดันข้อความลงมาหน่อย
            doc_copy.add_paragraph() 

        footer_texts = [
            f"ร.ต..............................ร่าง.......................{short_date}",
            f"   ............................พิมพ์/ทาน....................{short_date}",
            f"ร.ท.............................ตรวจ.......................{short_date}"
        ]

        for text in footer_texts:
            p_foot = doc_copy.add_paragraph()
            p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT # ดันชิดขวาสุดของหน้ากระดาษ
            run_foot = p_foot.add_run(text)
            run_foot.font.name = 'TH SarabunPSK'
            run_foot.font.size = Pt(16)

        doc_copy.save(dest_path)

    # สั่งทำสำเนาคู่ฉบับทั้ง 2 ไฟล์
    make_copy_doc(out_letter_path, out_copy_letter_path, "letter")
    make_copy_doc(out_memo_path, out_copy_memo_path, "memo")

    log_func(
        "\n🎉 สร้างเอกสารเสร็จสมบูรณ์เรียบร้อย!\n 📄 1) หนังสือภายนอก (ต้นฉบับ)\n 📄 2) บันทึกข้อความ (ต้นฉบับ)\n 📄 3) สำเนาคู่ฉบับหนังสือภายนอก\n 📄 4) สำเนาคู่ฉบับบันทึกข้อความ\n"
    )
    finish_callback(
        True,
        f"สร้างเอกสารสำหรับ {unique_id} ครบทั้ง 4 ฉบับ (บรรจุใน ZIP) เรียบร้อยแล้ว!",
        out_letter_path,
        out_memo_path,
        out_copy_letter_path,
        out_copy_memo_path
    )

  except Exception as e:
    log_func(f"\n❌ เกิดข้อผิดพลาด: {type(e).__name__} - {e}")
    finish_callback(False, str(e), None, None, None, None)


# =========================================================================
# 🎨 STREAMLIT WEB INTERFACE (ควบคุมการทำงานระบบออนไลน์)
# =========================================================================

# ดึงข้อมูลรหัส LP จาก Google Sheet พร้อม Caching
@st.cache_data(ttl=60)
def fetch_lps():
  scopes = [
      "https://www.googleapis.com/auth/spreadsheets",
      "https://www.googleapis.com/auth/drive",
  ]
  gc = gspread.service_account_from_dict(CREDENTIALS_DICT, scopes=scopes)
  ss = gc.open_by_key(SPREADSHEET_ID)
  ws = ss.worksheet("PrintLogs")
  data = ws.get_all_values()
  lps = [row[1].strip() for row in data[1:] if len(row) > 1 and row[1].strip()]
  return list(dict.fromkeys(reversed(lps)))[:30]


# 1. การ์ดตั้งค่าตัวเลือกสร้างเอกสาร
st.markdown(
    '<div style="background:#ffffff; padding:22px; border-radius:16px;'
    ' border:1px solid #e2e8f0; box-shadow:0 4px 12px'
    ' rgba(0,0,0,0.03);margin-bottom:20px;">',
    unsafe_allow_html=True,
)

template_mode = st.radio(
    "🌐 เลือกแหล่งที่มาของแม่แบบเอกสาร:",
    ["online", "local"],
    format_func=lambda x: (
        "ระบบออนไลน์ (OneDrive)"
        if x == "online"
        else "ระบบเครื่อง (.docx ในเครื่อง)"
    ),
    horizontal=True,
)

try:
  lp_options = fetch_lps()
  selected_lp = st.selectbox("📌 เลือกรหัส LP หรือระบุรหัสเอกสาร:", lp_options)
except Exception as e:
  st.error(f"ไม่สามารถเชื่อมต่อ Google Sheet ได้: {e}")
  selected_lp = st.text_input("กรอกรหัส LP เอง:")

st.markdown("</div>", unsafe_allow_html=True)

# =========================================================================
# 📌 ปุ่มสั่งสร้างเอกสาร Word พร้อมระบบล็อกรหัสผ่าน
# =========================================================================
SECRET_PASSWORD = "36529" 

doc_password = st.text_input(
    "🔒 กรอกรหัสผ่านเพื่ออนุมัติการสร้างเอกสาร:", type="password"
)

if st.button("🚀 เริ่มสร้างเอกสาร Word", type="primary"):
  if not selected_lp:
    st.warning("กรุณาเลือกรหัส LP ก่อนครับ")
  elif doc_password != SECRET_PASSWORD:
    st.error("❌ รหัสผ่านไม่ถูกต้อง! ไม่ได้รับอนุญาตให้สร้างเอกสารราชการฉบับนี้")
  else:
    log_area = st.empty()
    logs_list = []

    def web_log(text):
      logs_list.append(text)
      log_area.code("\n".join(logs_list), language="bash")

    def web_finish(success, message, letter_p, memo_p, copy_letter_p, copy_memo_p):
      if success:
        st.success(f"✅ {message}")
        
        # มัดรวมไฟล์ทั้งหมดลง ZIP ในหน่วยความจำ
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            if letter_p and os.path.exists(letter_p):
                zip_file.write(letter_p, f"๑_หนังสือภายนอก_{clean_lp}.docx")
            if memo_p and os.path.exists(memo_p):
                zip_file.write(memo_p, f"๒_บันทึกข้อความ(ต้นฉบับ)_{clean_lp}.docx")
            if copy_letter_p and os.path.exists(copy_letter_p):
                zip_file.write(copy_letter_p, f"๓_สำเนาหนังสือภายนอก_{clean_lp}.docx")
            if copy_memo_p and os.path.exists(copy_memo_p):
                zip_file.write(copy_memo_p, f"๔_สำเนาบันทึกข้อความ_{clean_lp}.docx")
        
        # ปุ่มดาวน์โหลดไฟล์ ZIP 
        st.download_button(
            label="📥 ดาวน์โหลดเอกสารทั้งหมด (ไฟล์ ZIP)",
            data=zip_buffer.getvalue(),
            file_name=f"Documents_{clean_lp}.zip",
            mime="application/zip"
        )
      else:
        st.error(f"❌ เกิดข้อผิดพลาด: {message}")

    clean_lp = selected_lp.replace(" (เคยสร้างแล้ว)", "")
    with st.spinner("กำลังประมวลผลสร้างเอกสาร..."):
      generate_documents_process(
          clean_lp, template_mode, web_log, web_finish
      )

# =========================================================================
# 📌 FOOTER SECTION (คณะผู้จัดทำ & ผู้ดูแลระบบ ถอดแบบต้นฉบับ 100%)
# =========================================================================
st.markdown("---")
st.markdown(
    """
<div class="custom-footer">
<div class="footer-system-title">
ระบบบริหารจัดการเอกสารและติดตามงบประมาณจ้างเหมาอัจฉริยะ (ฝ่ายแจ้งความต้องการ)
</div>

<div class="footer-leader">
<span class="material-icons" style="vertical-align: middle; color: #43a047; font-size: 1.2rem;">star</span>
อำนวยการและสนับสนุนการพัฒนาระบบโดย: <strong>น.อ. อมรพงศ์ เอี่ยมสะอาด</strong>
<br><small style="color: #616161;">(ผู้อำนวยการกองพัสดุช่างอากาศ)</small>
</div>

<div class="advisor-row">
<div class="advisor-col">
<span class="material-icons" style="vertical-align: middle; color: #fb8c00; font-size: 1.1rem;">lightbulb</span>
ที่ปรึกษาด้านข้อมูลและระบบ: <strong>ร.ท.ตุนท์ นามตาปี บำรุงศักดิ์</strong>
<br><small style="color: #757575;">(รอง หน.ฝจก.ผคค.กพอ.ชอ.)</small>
</div>
<div class="advisor-col">
<span class="material-icons" style="vertical-align: middle; color: #fb8c00; font-size: 1.1rem;">lightbulb</span>
ที่ปรึกษาด้านข้อมูลและระบบ: <strong>ร.ต.อาทิตย์ ศรีประสิทธิ์</strong>
<br><small style="color: #757575;">(หน.มว.แจ้งความต้องการต่างประเทศ ฝจก.ผคค.กพอ.ชอ.)</small>
</div>
</div>

<div class="dev-card">
<div style="font-size: 0.94rem; color: #37474f;">
<span class="material-icons" style="vertical-align: middle; color: #1e88e5; font-size: 1.2rem;">code</span>
ผู้พัฒนาและผู้ดูแลระบบ: <strong>ธรรศ วรวัฒนานุกูล</strong>
<br><small style="color: #616161;">(พนักงานบริการพัสดุ ฝจก.ผคค.กพอ.ชอ.)</small>
</div>
<div>
<a href="mailto:req-daesupply@requirements-asd.com?subject=_WEB%20System%20Error%20Report" class="btn-report">
<span class="material-icons" style="vertical-align: middle; font-size: 1.1rem; margin-right: 4px;">contact_support</span>
คลิกเพื่อแจ้งปัญหาการใช้งาน
</a>
</div>
</div>

<div class="copyright-text">
© 2026 ASD REQUIREMENTS SUPPLY | All Rights Reserved
</div>
</div>
""",
unsafe_allow_html=True,
)
